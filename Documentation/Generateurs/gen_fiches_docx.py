# -*- coding: utf-8 -*-
"""Genere la fiche d'exercice au format Word, illustree, pour chaque exercice.

Deux documents par exercice, sur le modele des .gh et des fiches Markdown :
    <ID>_fiche.docx         sujet ET corrige, illustres -> formateur
    <ID>_fiche_sujet.docx   sujet seul, illustre        -> apprenant

Les illustrations sont les captures de canvas produites par
Documentation/Generateurs/GH/gen_images.py.

Dependances : python-docx, Pillow.
"""
import os
import sys

ICI = os.path.dirname(os.path.abspath(__file__))
for p in (ICI, os.path.join(ICI, "GH")):
    if p not in sys.path:
        sys.path.insert(0, p)

from lots import TOUS as LOT_A, LOTS as _REGISTRE
from skill_a import SKILL

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

PROJET = os.path.abspath(os.path.join(ICI, "..", ".."))
SORTIES = [os.path.join(PROJET, d.replace("/", os.sep))
           for _c, _n, d, _l in _REGISTRE]
SORTIE = SORTIES[0]

VERSION = "v0.3-260826"
INDICE = "Ind. B"
DATE = "26/08/2026"
LOT = u"A — Découverte des composants natifs"

BLEU = RGBColor(0x2E, 0x5C, 0x8A)
GRIS = RGBColor(0x44, 0x4A, 0x52)
ROUGE = RGBColor(0xB0, 0x00, 0x20)

LARGEUR_MAX = 16.4          # cm, largeur utile de la page A4 marges comprises
HAUTEUR_MAX = 19.0          # cm : au-dela le canvas devient illisible dans la fiche

MANUELS = {}
try:
    import recipes_a1, recipes_a2, recipes_a3, recipes_a4
    for m in (recipes_a1, recipes_a2, recipes_a3, recipes_a4):
        for k, v in m.R.items():
            if v.get("manuel"):
                MANUELS[k] = v["manuel"]
except Exception as e:
    print("  (réglages manuels non repris : %s)" % e)


# ------------------------------------------------------------------- outillage
def ombrer(cellule, couleur):
    tcPr = cellule._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), couleur)
    tcPr.append(shd)


def bordures(cellule, couleur="BFC7D1"):
    tcPr = cellule._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for cote in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + cote)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), couleur)
        borders.append(el)
    tcPr.append(borders)


def para(doc, texte, taille=10.5, gras=False, couleur=None, avant=0, apres=4,
         italique=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(avant)
    p.paragraph_format.space_after = Pt(apres)
    r = p.add_run(texte)
    r.font.size = Pt(taille)
    r.font.bold = gras
    r.font.italic = italique
    r.font.name = "Calibri"
    if couleur is not None:
        r.font.color.rgb = couleur
    return p


def titre_section(doc, texte, couleur=BLEU):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texte)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = couleur
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), "8")
    bas.set(qn("w:space"), "3")
    bas.set(qn("w:color"), "%02X%02X%02X" % (couleur[0], couleur[1], couleur[2])
            if isinstance(couleur, tuple) else "2E5C8A")
    bdr.append(bas)
    pPr.append(bdr)
    return p


def sous_titre(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(texte)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = GRIS
    return p


def encadre(doc, texte, fond="EAF1F8", couleur=None, taille=11):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    c.width = Cm(LARGEUR_MAX)
    ombrer(c, fond)
    bordures(c)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texte)
    r.font.size = Pt(taille)
    r.font.name = "Calibri"
    if couleur is not None:
        r.font.color.rgb = couleur
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def puces(doc, elements, marqueur=u"•  "):
    for e in elements:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(marqueur + e)
        r.font.size = Pt(10.5)
        r.font.name = "Calibri"


def etapes_numerotees(doc, elements):
    for i, e in enumerate(elements, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(u"Étape %d.  " % i)
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.name = "Calibri"
        r.font.color.rgb = BLEU
        r2 = p.add_run(e)
        r2.font.size = Pt(10.5)
        r2.font.name = "Calibri"


def image(doc, chemin, legende=None):
    if not os.path.isfile(chemin):
        return False
    try:
        w, h = Image.open(chemin).size
    except Exception:
        w, h = (1600, 1000)
    larg = LARGEUR_MAX
    haut = larg * h / float(w)
    if haut > HAUTEUR_MAX:
        haut = HAUTEUR_MAX
        larg = haut * w / float(h)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(chemin, width=Cm(larg))
    if legende:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)
        r = c.add_run(legende)
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.name = "Calibri"
        r.font.color.rgb = GRIS
    return True


def tableau_infos(doc, e):
    lignes = [
        (u"Thématique", e["them"]),
        (u"Référence au référentiel", e["ref"]),
    ]
    if e.get("competence") and e["competence"] != u"—":
        lignes.append((u"Compétence visée", e["competence"]))
    if e.get("bloom"):
        lignes.append((u"Case Bloom (révisée)", e["bloom"]))
    lignes += [
        (u"Niveau", e["niv"]),
        (u"Durée cible", u"%d minutes" % e["duree"]),
        (u"Prérequis", e["prereq"]),
    ]
    if e.get("verdict") == u"connaissance":
        lignes.append((u"Mode de validation", u"— (non notée)"))
    else:
        lignes.append((u"Mode de validation",
                       u"%s — tolérance %s" % (e["mode"], e["tol"])))
        lignes.append((u"Solution de référence", u"%s composants" % e["nb"]))
    lignes += [
        (u"Version", u"%s — %s — %s" % (VERSION, INDICE, DATE)),
        (u"Conception", SKILL),
    ]
    t = doc.add_table(rows=len(lignes), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(lignes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        c0.width = Cm(5.2)
        c1.width = Cm(LARGEUR_MAX - 5.2)
        ombrer(c0, "EDF1F5")
        bordures(c0)
        bordures(c1)
        for cel, txt, gras in ((c0, k, True), (c1, v, False)):
            p = cel.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(txt)
            r.font.size = Pt(9.5)
            r.font.bold = gras
            r.font.name = "Calibri"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


# -------------------------------------------------------------------- document
def nouveau(e):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    for a in ("left_margin", "right_margin"):
        setattr(s, a, Cm(2.3))
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(u"%s — %s" % (e["id"], e["titre"]))
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = BLEU

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(u"Fiche d'exercice Magpie  ·  Lot %s" % LOT)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.name = "Calibri"
    r.font.color.rgb = GRIS

    # pied de page
    pied = s.footer.paragraphs[0]
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pied.add_run(u"Magpie · %s — %s · %s %s" % (e["id"], e["titre"], VERSION, INDICE))
    rr.font.size = Pt(8)
    rr.font.color.rgb = GRIS
    rr.font.name = "Calibri"
    return doc


def bloc_sujet(doc, e, dossier, seul):
    titre_section(doc, u"SUJET")

    sous_titre(doc, u"Compétence visée")
    para(doc, e["obj"])

    if e.get("contexte"):
        sous_titre(doc, u"Contexte")
        para(doc, e["contexte"])

    sous_titre(doc, u"Énoncé")
    encadre(doc, e["enonce"], fond="EAF1F8", taille=11)

    ill = os.path.join(dossier, "Illustrations", "%s_canvas_sujet.png" % e["id"])
    if image(doc, ill, u"Le canvas à l'ouverture de %s_sujet.gh : "
                       u"les données fournies et le paramètre REPONSE." % e["id"]):
        pass

    sous_titre(doc, u"Ce qui vous est fourni")
    para(doc, e["depart"])

    sous_titre(doc, u"Ce qui est attendu")
    para(doc, e["att"])
    para(doc, u"Branchez votre résultat sur le paramètre REPONSE. La correction compare "
              u"cette sortie en mode %s%s."
         % (e["mode"], "" if e["tol"] in (u"—", "0")
            else u", avec une tolérance de %s" % e["tol"]), italique=True)

    para(doc, u"La consigne ne nomme aucun composant, et c'est délibéré : nommer "
              u"l'outil reviendrait à donner la réponse. Ce lot n'autorise que des "
              u"composants natifs de Grasshopper pour Rhino 8 — aucun plugin tiers "
              u"n'est nécessaire.", taille=9.5, italique=True)

    sous_titre(doc, u"Barème")
    para(doc, e["bareme"])

    if seul:
        para(doc, u"Le corrigé fait l'objet d'une fiche distincte, remise après validation "
                  u"ou en fin de séance.", italique=True, avant=10)


def bloc_corrige(doc, e, dossier):
    doc.add_page_break()
    titre_section(doc, u"CORRIGÉ")

    encadre(doc, u"À ne consulter qu'après avoir cherché. Dans le fichier %s_complet.gh, "
                 u"le corrigé occupe la zone basse du canvas. Il ne produit rien tant que "
                 u"l'interrupteur AFFICHER LE CORRIGÉ n'est pas basculé sur vrai : "
                 u"positionnez-le sur faux pour faire disparaître le résultat."
                 % e["id"], fond="FDF3D6", taille=10)

    ill = os.path.join(dossier, "Illustrations", "%s_canvas_corrige.png" % e["id"])
    image(doc, ill, u"La zone corrigé. Elle est autonome : les données fournies y sont "
                    u"recopiées, aucun câble ne la relie à la zone sujet.")

    sous_titre(doc, u"Marche à suivre")
    etapes_numerotees(doc, e["etapes"])

    if e.get("erreur"):
        sous_titre(doc, u"L'erreur attendue")
        para(doc, u"C'est l'erreur qu'il faut guetter, parce qu'elle est diagnostique : "
                  u"elle dit ce que l'apprenant a mal compris, là où un simple « faux » "
                  u"ne dirait rien.", taille=9.5, italique=True)
        encadre(doc, e["erreur"], fond="FBEAEA", taille=10)

    sous_titre(doc, u"Pièges fréquents")
    puces(doc, e["pieges"])

    sous_titre(doc, u"Composants de la solution de référence")
    para(doc, e["comp"])
    para(doc, u"Cette liste figure sur la fiche formateur uniquement : elle ne doit "
              u"pas être remise à l'apprenant avant qu'il ait cherché.",
         taille=9.5, italique=True)

    if e.get("donnees_note"):
        sous_titre(doc, u"Pourquoi ce jeu de données")
        para(doc, e["donnees_note"])

    if e.get("limite"):
        sous_titre(doc, u"Limite de la correction automatique")
        encadre(doc, e["limite"], fond="FDF3D6", taille=10)

    if e.get("alerte"):
        sous_titre(doc, u"Note au formateur")
        encadre(doc, e["alerte"], fond="EFEAF6", taille=10)

    if e["id"] in MANUELS:
        sous_titre(doc, u"Réglages à poser à la main")
        para(doc, u"Ces réglages ne peuvent pas être enregistrés dans le fichier : "
                  u"ils sont à poser dans Grasshopper.", taille=9.5, italique=True)
        puces(doc, MANUELS[e["id"]])

    sous_titre(doc, u"Pour aller plus loin")
    puces(doc, e["var"])

    sous_titre(doc, u"Fichiers de cet exercice")
    fichiers = [
        (u"%s_sujet.gh" % e["id"], u"énoncé et données de départ, sans le corrigé"),
        (u"%s_complet.gh" % e["id"], u"énoncé et corrigé, ce dernier commandé par l'interrupteur"),
        (u"%s.json" % e["id"], u"descripteur pour le plugin Magpie"),
        (u"%s_fiche.docx" % e["id"], u"la présente fiche"),
        (u"%s_fiche_sujet.docx" % e["id"], u"la fiche sans le corrigé, pour l'apprenant"),
    ]
    rd = os.path.join(dossier, "Ressources")
    if os.path.isdir(rd):
        for f in os.listdir(rd):
            if f.lower().endswith(".3dm"):
                fichiers.append((u"Ressources/" + f,
                                 u"géométrie Rhino à ouvrir avant de commencer"))
    puces(doc, [u"%s — %s" % (a, b) for a, b in fichiers])



def bloc_charniere(doc, e, dossier):
    """Fiche d'un item requalifie en connaissance : pas un exercice note."""
    titre_section(doc, u"POURQUOI CE N'EST PAS UN EXERCICE")

    encadre(doc, u"Cet item n'est pas un exercice noté. Il porte une connaissance "
                 u"nécessaire, mais qui s'acquiert et se vérifie par une question, "
                 u"non par un montage : la construire dans Grasshopper mesurerait "
                 u"la mémoire, pas la compétence.", fond="FDF3D6", taille=10.5)

    para(doc, u"L'énoncé d'origine demandait de constater un comportement du logiciel "
              u"plutôt que de produire un résultat. La réponse s'obtenait en sachant, "
              u"non en construisant.", avant=6)

    sous_titre(doc, u"Énoncé d'origine, conservé pour mémoire")
    para(doc, e.get("enonce_origine", e["enonce"]), italique=True)

    if e.get("contexte"):
        sous_titre(doc, u"Contexte")
        para(doc, e["contexte"])

    titre_section(doc, u"LA QUESTION")
    for ligne in (e.get("charniere") or u"").split(chr(10)):
        if ligne.strip():
            gras = ligne.strip().startswith(u"Valeur diagnostique")
            para(doc, ligne, gras=gras)

    titre_section(doc, u"COMMENT L'EMPLOYER")
    puces(doc, [
        u"Avant l'exercice qui mobilise cette connaissance, pas après : "
        u"elle en est un prérequis.",
        u"Poser la question à main levée, relever la répartition des réponses, "
        u"et n'expliquer que si une réponse fausse est majoritaire.",
        u"La valeur est dans la mauvaise réponse : elle nomme la représentation "
        u"à corriger.",
    ])

    ill = os.path.join(dossier, "Illustrations", "%s_canvas_corrige.png" % e["id"])
    if os.path.isfile(ill):
        sous_titre(doc, u"Démonstration facultative")
        para(doc, u"Le fichier %s_complet.gh reste disponible comme support de "
                  u"démonstration au vidéoprojecteur. Il n'est pas à faire "
                  u"construire." % e["id"], taille=9.5, italique=True)
        image(doc, ill, u"Le montage de démonstration.")


def main():
    dossiers = {}
    for _s in SORTIES:
        if not os.path.isdir(_s):
            continue
        for d in os.listdir(_s):
            if os.path.isdir(os.path.join(_s, d)):
                dossiers[d.split(" ")[0]] = os.path.join(_s, d)
    if not dossiers:
        print("Aucun dossier d'exercice trouve.")
        return

    faits, charnieres, absents, sans_image = 0, 0, [], []
    verrouilles = []

    def _sauver(doc, chemin):
        try:
            doc.save(chemin)
            return True
        except PermissionError:
            verrouilles.append(os.path.basename(chemin))
            return False
    for e in LOT_A:
        dd = dossiers.get(e["id"])
        if dd is None:
            absents.append(e["id"])
            continue
        ill = os.path.join(dd, "Illustrations")
        if not os.path.isfile(os.path.join(ill, "%s_canvas_sujet.png" % e["id"])):
            sans_image.append(e["id"])

        if e.get("verdict") == u"connaissance":
            for nom in ("%s_fiche.docx", "%s_fiche_sujet.docx"):
                d = nouveau(e)
                tableau_infos(d, e)
                bloc_charniere(d, e, dd)
                _sauver(d, os.path.join(dd, nom % e["id"]))
            charnieres += 1
            continue

        d1 = nouveau(e)
        tableau_infos(d1, e)
        bloc_sujet(d1, e, dd, False)
        bloc_corrige(d1, e, dd)
        _sauver(d1, os.path.join(dd, "%s_fiche.docx" % e["id"]))

        d2 = nouveau(e)
        tableau_infos(d2, e)
        bloc_sujet(d2, e, dd, True)
        _sauver(d2, os.path.join(dd, "%s_fiche_sujet.docx" % e["id"]))
        faits += 1

    print("Fiches Word d'exercice        : %d" % faits)
    print("Fiches Word question charnière : %d" % charnieres)
    print("Documents écrits               : %d" % (2 * (faits + charnieres)))
    if verrouilles:
        print("NON ECRITS (fichier verrouille, probablement ouvert dans Word) : %s"
              % ", ".join(verrouilles))
    if sans_image:
        print("Sans illustration : %s" % ", ".join(sans_image))
    if absents:
        print("Sans dossier : %s" % ", ".join(absents))
    print("Lots traités : %d" % len(SORTIES))


if __name__ == "__main__":
    main()
